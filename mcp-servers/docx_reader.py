# /// script
# requires-python = ">=3.10"
# dependencies = ["mcp[cli]", "python-docx", "mammoth"]
# ///
"""
MCP Server: docx-reader
读取 .docx 文件，支持纯文本、Markdown、元数据提取
"""

import os
import json
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("docx-reader")

@mcp.tool()
def read_docx(file_path: str) -> str:
    """Read a .docx file and return plain text."""
    p = os.path.abspath(file_path)
    if not os.path.exists(p):
        return f"Error: not found: {p}"
    from docx import Document
    doc = Document(p)
    out = []
    for para in doc.paragraphs:
        if para.text.strip():
            out.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append(" | ".join(cells))
    return chr(10).join(out)

@mcp.tool()
def read_docx_as_markdown(file_path: str) -> str:
    """Read a .docx file and convert to Markdown."""
    p = os.path.abspath(file_path)
    if not os.path.exists(p):
        return f"Error: not found: {p}"
    import mammoth
    with open(p, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    return result.value

@mcp.tool()
def read_docx_info(file_path: str) -> str:
    """Get .docx metadata: title, author, paragraph and table count."""
    p = os.path.abspath(file_path)
    if not os.path.exists(p):
        return f"Error: not found: {p}"
    from docx import Document
    doc = Document(p)
    pr = doc.core_properties
    info = {
        "file": p,
        "title": pr.title or "",
        "author": pr.author or "",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "chars": sum(len(x.text) for x in doc.paragraphs)
    }
    return json.dumps(info, ensure_ascii=False, indent=2)

if __name__ == "__main__":
    mcp.run()
